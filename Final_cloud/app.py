import os
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
import logging
import boto3
import pypandoc
import time
import psutil
from pdf2docx import parse as pdf_to_docx
from docx import Document
import markdown
import html2text
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from botocore.exceptions import ClientError

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['ALLOWED_INPUT_FORMATS'] = {'txt', 'doc', 'docx', 'md', 'html', 'odt', 'rtf', 'pdf'}
app.config['ALLOWED_OUTPUT_FORMATS'] = {'doc', 'docx', 'pdf', 'md', 'html', 'odt', 'rtf', 'epub'}

# Configure logging
logging.basicConfig(filename='performance.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_INPUT_FORMATS']

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify(error='No file part'), 400

    file = request.files['file']
    input_format = request.form.get('input_format')
    output_format = request.form.get('output_format')

    if not input_format or not output_format:
        return jsonify(error='Input and output formats must be specified'), 400

    if file.filename == '':
        return jsonify(error='No selected file'), 400

    if file and allowed_file(file.filename):
        try:
            start_time = time.time()
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)

            filename = secure_filename(file.filename)
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            output_filename = secure_filename(file.filename.rsplit('.', 1)[0] + '.' + output_format)
            output_path = os.path.join(app.config['CONVERTED_FOLDER'], output_filename)

            # Perform conversion based on input and output formats
            if input_format == 'txt' and output_format in ['doc', 'docx']:
                txt_to_docx(input_path, output_path)
            elif input_format == 'docx' and output_format == 'txt':
                docx_to_txt(input_path, output_path)
            elif input_format == 'md' and output_format == 'html':
                md_to_html(input_path, output_path)
            elif input_format == 'html' and output_format == 'md':
                html_to_md(input_path, output_path)
            elif input_format == 'odt' and output_format in ['doc', 'docx']:
                odt_to_docx(input_path, output_path)
            elif input_format == 'rtf' and output_format in ['doc', 'docx']:
                rtf_to_docx(input_path, output_path)
            elif input_format == 'pdf' and output_format in ['doc', 'docx']:
                pdf_to_docx(input_path, output_path)
            else:
                # Use pypandoc for other conversions
                pypandoc.convert_file(input_path, output_format, outputfile=output_path)

            # Log conversion time
            end_time = time.time()
            conversion_time = end_time - start_time
            logging.info(f"Conversion time: {conversion_time} seconds")

            # Log conversion time
            end_time = time.time()
            conversion_time = end_time - start_time
            logging.info(f"Conversion time: {conversion_time} seconds")

            # Log CPU usage
            cpu_usage = psutil.cpu_percent(interval=None)
            logging.info(f"CPU usage: {cpu_usage}%")

            # Log memory usage
            memory_usage = psutil.virtual_memory().percent
            logging.info(f"Memory usage: {memory_usage}%")

            # AWS credentials configuration
            s3 = boto3.client(
                's3',
                aws_access_key_id='AKIA4MTWL7VCLFPZDT7H',
                aws_secret_access_key='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',
                region_name='us-east-1'  # Ensure this is the correct region
            )
            
            with open(output_path, 'rb') as f:
                s3.upload_fileobj(f, 'document-conversion-buckett', output_filename)

            return jsonify(message='File converted successfully', output_file=output_filename, conversion_time=conversion_time, cpu_usage=cpu_usage, memory_usage=memory_usage ), 200
        except Exception as e:
            # Log error
            logging.error(f"Error occurred: {str(e)}")
            return jsonify(error=str(e)), 500
    else:
        return jsonify(error='File not allowed'), 400

def txt_to_docx(txt_path, docx_path):
    with open(txt_path, 'r') as txt_file:
        txt_content = txt_file.read()
    doc = Document()
    doc.add_paragraph(txt_content)
    doc.save(docx_path)

def docx_to_txt(docx_path, txt_path):
    doc = Document(docx_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    with open(txt_path, 'w') as txt_file:
        txt_file.write(text)

def md_to_html(md_path, html_path):
    with open(md_path, 'r') as md_file:
        md_content = md_file.read()
    html_content = markdown.markdown(md_content)
    with open(html_path, 'w') as html_file:
        html_file.write(html_content)

def html_to_md(html_path, md_path):
    with open(html_path, 'r') as html_file:
        html_content = html_file.read()
    md_content = html2text.html2text(html_content)
    with open(md_path, 'w') as md_file:
        md_file.write(md_content)

def odt_to_docx(odt_path, docx_path):
    pypandoc.convert_file(odt_path, 'docx', outputfile=docx_path)

def rtf_to_docx(rtf_path, docx_path):
    # RTF to DOCX conversion logic goes here
    pass

def send_email(file_path):
    # Email configuration
    sender_email = "srinithyap21@gmail.com"
    receiver_email = "karivedatrisha@gmail.com"
    subject = "File Conversion Report"
    body = "Please find the attached file converted successfully."

     # AWS SES configuration
    aws_region = "us-east-1"  # Change to your AWS region

    #smtp_username = "AKIA4MTWL7VCKZNQNTN5"  # Generated SMTP username
    #smtp_password = "BFSG6+7/wgA7yKEt9Eami0UaUkoJm/DgX8eCXgqI7LqI"  # Generated SMTP password

    # Create a new SES resource and specify a region
    client = boto3.client('ses', region_name= "us-east-1")

# Try to send the email
    try:
        # Provide the contents of the email
        response = client.send_raw_email(
            Source=sender_email,
            Destinations=[receiver_email],
            RawMessage={
                'Data': create_mime_message(sender_email, receiver_email, subject, body, file_path)
            }
        )
        print("Email sent! Message ID:", response['MessageId'])
    except ClientError as e:
        print("Email sending failed:", e.response['Error']['Message'])

def create_mime_message(sender_email, receiver_email, subject, body, file_path):
    # Create a multipart message
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = subject

    # Add body to email
    msg.attach(MIMEText(body, 'plain'))

    # Add attachment
    attachment = MIMEApplication(open(file_path, 'rb').read())
    attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(file_path))
    msg.attach(attachment)

    # Return the MIME message as a string
    return msg.as_string()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
